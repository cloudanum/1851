"""build_ig.py — rebuild docs/1851-IGa2.docx for course 1851 revision A.2.

Idempotent: restores a byte copy of the A.1 skeleton
(FROM_ftp/1851a1/1851-IGa1.docx), then
  * Course History: keeps the A.1 revision line + description, and inserts a
    new A.2 revision block (bold revision tag line + description) cloned from
    the A.1 paragraphs so style/formatting match. The "Your feedback ..."
    callout is left untouched — its apparent duplication in document.xml is
    one shape stored as mc:AlternateContent (wps Choice + VML Fallback), i.e.
    a single rendered object.
  * Timeline: replaces the unfilled placeholder rows ('Chapter ', '##:##am',
    '# min') with the full A.2 3-day plan. Rows are deep copies of the
    skeleton's own header / Day / chapter / sub-row XML, so borders, shading,
    widths and the Title1 paragraph style are preserved.
  * Title page: no revision marker exists in the skeleton (verified: 'A.1'
    occurs only in the Course History line), so it is left untouched.

Run (from 1851a2-author-input/):
    /Users/iahmad/Creator/Courses_and_conferences/LT/.venv-courseware/bin/python tools/build_ig.py
"""
from __future__ import annotations

import copy
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT.parent / "FROM_ftp" / "1851a1" / "1851-IGa1.docx"
OUT = ROOT / "docs" / "1851-IGa2.docx"

# --- A.2 Course History --------------------------------------------------------
A2_DESC = (
    "Major revision. The course was modernized to be GenAI-first: prompt "
    "engineering with the OpenAI API was expanded into its own chapter; new "
    "chapters were added on generative AI across the software development "
    "lifecycle (including test-driven development with AI), evaluating "
    "generative AI systems, and AI agents and agentic workflows; traditional "
    "machine-learning content was condensed into a single Machine Learning "
    "Essentials chapter; and security coverage was extended to agentic-era "
    "threats. New hands-on labs: 1.1, 2.1, 3.1, 4.1, 5.1, 6.1, 7.1, 8.1, "
    "9.1 — three per day. New activities: 1.2, 4.2. Do Now warm-up exercises "
    "were added throughout."
)

# --- A.2 timeline --------------------------------------------------------------
# 3-day course, 8:30am start, two 15-min breaks + 1-hour lunch per day.
# Lab rows sit immediately after their chapter block (Lab 5.1 and Lab 8.1 open
# the afternoon so no chapter/lab row spans lunch); Do Now warm-ups ride inside
# chapter lecture time and get no rows of their own.
# Day 1: 335 min, Day 2: 375 min, Day 3: 310 min (light final day).
TIMELINE = [
    ("day", "Day 1"),
    ("chapter", "Chapter 0", "Introduction and Overview", "8:30am", "15 min"),
    ("chapter", "Chapter 1", "Introduction to AI in Software Development", "8:45am", "70 min"),
    ("break", "9:55am", "15 min"),
    ("lab", "Lab 1.1", "Assistant Recon — Where AI Helps, Where It Fails", "10:10am", "30 min"),
    ("chapter", "Chapter 2", "Quality Characteristics and Ethics in AI", "10:40am", "45 min"),
    ("lab", "Lab 2.1", "Hallucination and Bias Audit", "11:25am", "30 min"),
    ("lunch", "11:55am", "60 min"),
    ("chapter", "Chapter 3", "Machine Learning Essentials", "12:55pm", "100 min"),
    ("break", "2:35pm", "15 min"),
    ("lab", "Lab 3.1", "Decision Trees vs. an LLM", "2:50pm", "45 min"),
    ("day", "Day 2"),
    ("chapter", "Chapter 4", "Prompt Engineering with OpenAI", "8:30am", "85 min"),
    ("break", "9:55am", "15 min"),
    ("lab", "Lab 4.1", "Building a PTCF Agent with the OpenAI API", "10:10am", "45 min"),
    ("chapter", "Chapter 5", "Generative AI Across the Software Development Lifecycle", "10:55am", "85 min"),
    ("lunch", "12:20pm", "60 min"),
    ("lab", "Lab 5.1", "TDD with an AI Pair", "1:20pm", "45 min"),
    ("break", "2:05pm", "15 min"),
    ("chapter", "Chapter 6", "Evaluating Generative AI Systems", "2:20pm", "70 min"),
    ("lab", "Lab 6.1", "Write an Eval Suite", "3:30pm", "45 min"),
    ("day", "Day 3"),
    ("chapter", "Chapter 7", "AI Agents and Agentic Workflows", "8:30am", "80 min"),
    ("break", "9:50am", "15 min"),
    ("lab", "Lab 7.1", "Build a Guarded Agent", "10:05am", "50 min"),
    ("chapter", "Chapter 8", "AI Security and Vulnerability Testing", "10:55am", "75 min"),
    ("lunch", "12:10pm", "60 min"),
    ("lab", "Lab 8.1", "Red Team the Agent", "1:10pm", "40 min"),
    ("chapter", "Chapter 9", "Course Summary", "1:50pm", "15 min"),
    ("break", "2:05pm", "15 min"),
    ("lab", "Lab 9.1", "Capstone — One Ticket, End to End", "2:20pm", "50 min"),
]


def set_cell_text(cell, text: str, bold: bool | None = None) -> None:
    """Replace a cell's text in place, keeping its paragraph style/format."""
    para = cell.paragraphs[0]
    # drop any extra paragraphs, keep the first
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r._r.getparent().remove(r._r)
        run = para.runs[0]
    else:
        run = para.add_run(text)
    if bold is not None:
        run.font.bold = bold


def edit_course_history(doc: Document) -> None:
    paras = doc.paragraphs
    rev_idx = next(
        i for i, p in enumerate(paras) if p.text.startswith("1851/CN/A.1/505/-")
    )
    a1_rev = paras[rev_idx]        # bold revision tag line
    a1_desc = paras[rev_idx + 1]   # description paragraph
    assert a1_desc.text == "This is the first public revision of this course."
    blank = next(p for p in paras[rev_idx + 2:] if p.style.name == "hist" and not p.text)

    # A.2 revision tag line — clone the A.1 line run by run
    a2_rev = copy.deepcopy(a1_rev._p)
    for r in a2_rev.findall(f"{'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'}r"):
        t = r.find(f"{'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'}t")
        if t is None:
            continue
        if t.text == "A.1":
            t.text = "A.2"
        elif t.text == "May":
            t.text = "August"
        elif t.text == "2025":
            t.text = "2026"

    # A.2 description — clone the A.1 description paragraph
    a2_desc = copy.deepcopy(a1_desc._p)
    a2_desc_para = Paragraph(a2_desc, a1_desc._parent)
    a2_desc_para.runs[0].text = A2_DESC

    # insert after the A.1 block: blank separator, tag line, description
    a1_desc._p.addnext(copy.deepcopy(blank._p))
    blank_new = a1_desc._p.getnext()
    blank_new.addnext(a2_rev)
    a2_rev.addnext(a2_desc)


def edit_timeline(doc: Document) -> None:
    tbl_el = next(
        t._tbl
        for t in doc.tables
        if [c.text for c in t.rows[0].cells] == ["", "Title", "Start Time", "Duration"]
    )
    table = Table(tbl_el, doc)

    # template rows from the skeleton (deep copies before we delete anything)
    day_tr = copy.deepcopy(table.rows[1]._tr)                     # shaded 'Day 1'
    ch_tr = copy.deepcopy(table.rows[2]._tr)                      # bold 'Chapter ' row
    sub_tr = copy.deepcopy(table.rows[3]._tr)                     # non-bold sub-row

    # drop every row except the header
    for tr in tbl_el.findall(f"{'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'}tr")[1:]:
        tbl_el.remove(tr)

    for entry in TIMELINE:
        kind = entry[0]
        if kind == "day":
            tr = copy.deepcopy(day_tr)
        elif kind in ("chapter", "lab"):
            tr = copy.deepcopy(ch_tr)
        else:
            tr = copy.deepcopy(sub_tr)
        tbl_el.append(tr)

    # fill texts through the python-docx API on the rebuilt table
    table = Table(tbl_el, doc)
    for row, entry in zip(table.rows[1:], TIMELINE):
        cells = row.cells
        if entry[0] == "day":
            set_cell_text(cells[0], entry[1])
        elif entry[0] in ("chapter", "lab"):
            _, label, title, start, dur = entry
            set_cell_text(cells[0], label)
            set_cell_text(cells[1], title)
            set_cell_text(cells[2], start)
            set_cell_text(cells[3], dur)
        else:  # lunch/break — non-bold sub-row (template run already carries b=0)
            _, start, dur = entry
            set_cell_text(cells[0], entry[0].capitalize(), bold=False)
            set_cell_text(cells[2], start, bold=False)
            set_cell_text(cells[3], dur, bold=False)


def validate(path: Path) -> None:
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    tables = [[[c.text for c in r.cells] for r in t.rows] for t in doc.tables]
    all_text = "\n".join(texts + [c for t in tables for r in t for c in r])

    # Course History
    assert all_text.count("1851/CN/A.1/505/-") == 1, "A.1 revision line missing/duplicated"
    assert all_text.count("1851/CN/A.2/505/-") == 1, "A.2 revision line missing/duplicated"
    assert "August 2026" in all_text
    assert A2_DESC in texts, "A.2 description paragraph not found verbatim"

    # Timeline — no placeholders left, fully populated
    for bad in ("##", "#:00", "# min", "Chapter \n"):
        assert bad not in all_text, f"placeholder {bad!r} still present"
    timeline = next(t for t in tables if t[0] == ["", "Title", "Start Time", "Duration"])
    assert len(timeline) == 1 + len(TIMELINE), f"timeline rows: {len(timeline)}"
    for row, entry in zip(timeline[1:], TIMELINE):
        if entry[0] == "day":
            assert row[0] == entry[1] and not row[1] and not row[2] and not row[3], row
        elif entry[0] in ("chapter", "lab"):
            assert row == [entry[1], entry[2], entry[3], entry[4]], row
        else:
            assert row == [entry[0].capitalize(), "", entry[1], entry[2]], row
    labs = [e[1] for e in TIMELINE if e[0] == "lab"]
    assert labs == [f"Lab {i}.1" for i in range(1, 10)], f"lab rows: {labs}"
    total = sum(int(e[4].split()[0]) for e in TIMELINE if e[0] in ("chapter", "lab"))
    assert total == (15 + 70 + 45 + 100 + 85 + 85 + 70 + 80 + 75 + 15) + (
        30 + 30 + 45 + 45 + 45 + 45 + 50 + 40 + 50
    ) == 1020

    # zip integrity
    assert zipfile.ZipFile(path).testzip() is None
    r = subprocess.run(["unzip", "-t", str(path)], capture_output=True, text=True)
    assert r.returncode == 0, r.stdout + r.stderr

    print("--- Course History paragraphs ---")
    hist_on = False
    for p in doc.paragraphs:
        if p.text == "Course History":
            hist_on = True
        if p.text == "Timeline":
            hist_on = False
        if hist_on and p.text.strip():
            print(f"[{p.style.name}] {p.text}")
    print("\n--- Timeline table ---")
    for row in timeline:
        print(" | ".join(row))
    day_totals: dict[str, int] = {}
    cur_day = ""
    for e in TIMELINE:
        if e[0] == "day":
            cur_day = e[1]
            day_totals[cur_day] = 0
        elif e[0] in ("chapter", "lab"):
            day_totals[cur_day] += int(e[4].split()[0])
    print("\nper-day teaching: " + ", ".join(f"{d}: {m} min" for d, m in day_totals.items()))
    print(f"teaching total: {total} min over 3 days; rows: {len(timeline)}")
    print("zip integrity: OK")


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(SRC, OUT)  # byte copy — idempotent rebuild
    doc = Document(OUT)
    edit_course_history(doc)
    edit_timeline(doc)
    doc.save(OUT)
    validate(OUT)
    print(f"\nbuilt {OUT}")


if __name__ == "__main__":
    sys.exit(main())
